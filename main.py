# ==============================================================================
# 1. 라이브러리 임포트
# ==============================================================================
import io
import os
import re
import requests
import docx
import json
import google.generativeai as genai
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_ORIENTATION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

# ==============================================================================
# 2. Flask 앱 초기화 및 설정
# ==============================================================================
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}},
     allow_headers=["Authorization", "Content-Type"],
     methods=["GET", "POST", "OPTIONS"],
     supports_credentials=True)

# --- Google Sheets 설정 ---
SHEET_API_SCOPES = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
# 환경 변수에서 서비스 계정 키(JSON 내용)를 직접 읽어옴
try:
    creds_json_str = os.environ.get("GOOGLE_SHEETS_CREDENTIALS")
    if not creds_json_str:
        raise ValueError("환경 변수 'GOOGLE_SHEETS_CREDENTIALS'가 설정되지 않았습니다.")
    creds_dict = json.loads(creds_json_str)
    SHEET_CREDS = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, SHEET_API_SCOPES)
    SHEET_CLIENT = gspread.authorize(SHEET_CREDS)
    print("Google Sheets API가 성공적으로 초기화되었습니다.")
except Exception as e:
    SHEET_CLIENT = None
    print(f"!!! Google Sheets API 초기화 오류: {e}")


# --- Gemini API 설정 ---
try:
    API_KEY = os.environ.get("GEMINI_API_KEY")
    if not API_KEY:
        raise ValueError("환경 변수 'GEMINI_API_KEY'가 설정되지 않았습니다.")

    genai.configure(api_key=API_KEY)
    model = genai.GenerativeModel('gemini-1.5-pro-latest')
    print("Gemini API 모델이 성공적으로 초기화되었습니다.")

except Exception as e:
    print(f"!!! Gemini API 초기화 오류: {e}")
    model = None

# ==============================================================================
# 3. 워드 문서 생성 헬퍼 함수 (변경 없음)
# ==============================================================================
def create_page_number_field(paragraph, field_text):
    run = paragraph.add_run()
    fldChar_begin = docx.oxml.shared.OxmlElement('w:fldChar'); fldChar_begin.set(docx.oxml.ns.qn('w:fldCharType'), 'begin'); run._r.append(fldChar_begin)
    run = paragraph.add_run(); instrText = docx.oxml.shared.OxmlElement('w:instrText'); instrText.set(docx.oxml.ns.qn('xml:space'), 'preserve'); instrText.text = field_text; run._r.append(instrText)
    run = paragraph.add_run(); fldChar_separate = docx.oxml.shared.OxmlElement('w:fldChar'); fldChar_separate.set(docx.oxml.ns.qn('w:fldCharType'), 'separate'); run._r.append(fldChar_separate)
    run = paragraph.add_run(); t = docx.oxml.shared.OxmlElement('w:t'); t.text = '1'; run._r.append(t)
    run = paragraph.add_run(); fldChar_end = docx.oxml.shared.OxmlElement('w:fldChar'); fldChar_end.set(docx.oxml.ns.qn('w:fldCharType'), 'end'); run._r.append(fldChar_end)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for border_name, border_attrs in kwargs.items():
        for border in tcBorders.findall(qn(f'w:{border_name}')):
            tcBorders.remove(border)
        border_element = OxmlElement(f'w:{border_name}')
        for attr_name, attr_val in border_attrs.items():
            border_element.set(qn(f'w:{attr_name}'), str(attr_val))
        tcBorders.append(border_element)

def add_image_border(run, border_width_pt=0, border_color='000000'):
    border_width_emu = int(border_width_pt * 12700)
    r = run._r
    drawing = r.find(qn('w:drawing'))
    if drawing is None: return
    try:
        spPr = drawing.xpath('.//pic:spPr')[0]
    except IndexError:
        return
    for ln in spPr.findall(qn('a:ln')):
        spPr.remove(ln)
    ln = OxmlElement('a:ln'); ln.set('w', str(border_width_emu)); ln.set('cap', 'flat'); ln.set('cmpd', 'sng')
    solidFill = OxmlElement('a:solidFill'); srgbClr = OxmlElement('a:srgbClr'); srgbClr.set('val', border_color); solidFill.append(srgbClr); ln.append(solidFill)
    prstDash = OxmlElement('a:prstDash'); prstDash.set('val', 'solid'); ln.append(prstDash)
    miter = OxmlElement('a:miter'); miter.set('lim', '800000'); ln.append(miter)
    spPr.append(ln)

def insert_image_to_paragraph(p, image_url_or_id, section):
    try:
        drive_match = re.search(r'/file/d/([a-zA-Z0-9_-]+)', image_url_or_id)
        image_url = f'https://drive.google.com/uc?export=download&id={drive_match.group(1)}' if drive_match else image_url_or_id
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(image_url, stream=True, headers=headers)
        response.raise_for_status()
        image_stream = io.BytesIO(response.content)
        with Image.open(image_stream) as img:
            native_width_px, _ = img.size
        max_width_emu = section.page_width - section.left_margin - section.right_margin
        if p.style.name.startswith('Table'): max_width_emu -= Cm(0.5)
        native_width_emu = native_width_px * 9525
        image_stream.seek(0)
        run = p.add_run()
        if native_width_emu > max_width_emu:
            run.add_picture(image_stream, width=max_width_emu)
        else:
            run.add_picture(image_stream)
        add_image_border(run)
    except Exception as e:
        p.add_run(f"[이미지 로드 오류: {e}]")

def merge_empty_cells_in_table(table):
    for row in table.rows:
        for i in range(len(row.cells) - 1, 0, -1):
            if row.cells[i].text.strip() == '':
                left_cell = row.cells[i-1]
                left_cell.merge(row.cells[i])

def parse_complex_table_data(raw_lines, num_cols=24):
    final_table_data = []
    for line in raw_lines:
        row_data = [''] * num_cols
        segments = line.split('|')
        current_parse_col = 0
        for segment in segments:
            match = re.search(r'^(.*?)\s*\{([\d]+|-)\}\s*$', segment)
            if match:
                text_content, col_indicator = match.group(1).strip(), match.group(2)
                target_col = num_cols - 1 if col_indicator == '-' else int(col_indicator) - 1
                if 0 <= target_col < num_cols:
                    row_data[target_col] = text_content
                    current_parse_col = target_col + 1
            else:
                if current_parse_col < num_cols:
                    row_data[current_parse_col] = segment
                    current_parse_col += 1
        final_table_data.append(row_data)
    return final_table_data

def generate_dynamic_filename(title):
    if not title:
        kst = timezone(timedelta(hours=9))
        now_kst = datetime.now(kst)
        date_str = now_kst.strftime('%y%m%d')
        time_str = now_kst.strftime('%H%M%S')
        filename = f"제목없음_{date_str}_T{time_str}.docx"
    else:
        filename = f"{title}.docx"
    return filename

def create_word_document(text_content, settings):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font_name = settings.get('font_family_east_asia', '맑은 고딕')
    font_size = settings.get('font_size', 10.5)
    line_spacing = settings.get('line_spacing', 1.0)
    para_spacing_after = settings.get('para_spacing_after', 8)
    font.name = font_name
    font.size = Pt(font_size)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    p_format = style.paragraph_format
    p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(para_spacing_after)
    p_format.line_spacing = line_spacing
    p_format.first_line_indent = Pt(0)
    pPr = style.element.pPr
    kinsoku = docx.oxml.shared.OxmlElement('w:kinsoku'); kinsoku.set(docx.oxml.ns.qn('w:val'), 'true'); pPr.append(kinsoku)
    wordWrap = docx.oxml.shared.OxmlElement('w:wordWrap'); wordWrap.set(docx.oxml.ns.qn('w:val'), 'false'); pPr.append(wordWrap)
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run('- '); create_page_number_field(footer_para, r'PAGE \* MERGEFORMAT'); footer_para.add_run(' -')
    for run in footer_para.runs:
        run.font.name = '바탕'; run.font.size = Pt(10); run._r.rPr.rFonts.set(qn('w:eastAsia'), '바탕')
    if settings.get('page_orientation') == 'LANDSCAPE':
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    else:
        section.orientation = WD_ORIENTATION.PORTRAIT
    section.top_margin = Cm(settings.get('margin_top', 2.0))
    section.bottom_margin = Cm(settings.get('margin_bottom', 2.0))
    section.left_margin = Cm(settings.get('margin_left', 2.5))
    section.right_margin = Cm(settings.get('margin_right', 2.5))

    lines = text_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip() == "{페이지바꿈}":
            doc.add_page_break()
            i += 1
            continue
        table_type = None
        if line.strip().startswith("{표시작1"):
            table_type = "simple"
        elif line.strip().startswith("{표시작2"):
            table_type = "complex"
        if table_type:
            table_settings = {}
            match = re.search(r'\{표시작\d([^}]*)\}', line)
            if match:
                params_str = match.group(1)
                if "테두리없음" in params_str: table_settings['borderless'] = True
                font_match = re.search(r'글꼴=([^,}]*)', params_str)
                if font_match: table_settings['font_name'] = font_match.group(1).strip()
                size_match = re.search(r'크기=([\d.]+)', params_str)
                if size_match: table_settings['font_size'] = float(size_match.group(1).strip())
            table_lines = []
            i += 1
            while i < len(lines) and not (lines[i].strip() == "{표끝1}" or lines[i].strip() == "{표끝2}"):
                table_lines.append(lines[i])
                i += 1
            table_data = []
            num_cols = 0
            if table_type == "simple":
                table_data = [l.split('|') for l in table_lines]
                if table_data:
                    num_cols = max(len(row) for row in table_data) if table_data else 0
                for row in table_data:
                    row.extend([''] * (num_cols - len(row)))
            elif table_type == "complex":
                num_cols = 24
                table_data = parse_complex_table_data(table_lines, num_cols)
            if len(table_data) > 0 and num_cols > 0:
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                table.style = 'Table Grid'
                for r_idx, row_data in enumerate(table_data):
                    row_cells = table.rows[r_idx].cells
                    first_cell_text = str(row_data[0]) if row_data else ""
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < len(row_cells):
                            cell = row_cells[c_idx]
                            cell_p = cell.paragraphs[0]
                            if not table_settings.get('borderless'):
                                cell_p.paragraph_format.space_after = Pt(0)
                                cell_p.paragraph_format.line_spacing = 1.5
                            cell_p.clear()
                            image_match_in_cell = re.search(r'\{그림:([^}]+)\}', cell_text.strip())
                            if image_match_in_cell:
                                image_url_or_id = image_match_in_cell.group(1).strip()
                                cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                cell.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
                                insert_image_to_paragraph(cell_p, image_url_or_id, section)
                            else:
                                run = cell_p.add_run(str(cell_text))
                                if not table_settings.get('borderless'):
                                    run.font.name = doc.styles['Normal'].font.name
                                    run.font.size = Pt(10)
                                    run._r.rPr.rFonts.set(qn('w:eastAsia'), doc.styles['Normal'].element.rPr.rFonts.get(qn('w:eastAsia')))
                    if "{회색}" in first_cell_text:
                        for cell in row_cells:
                            shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'D9D9D9'); cell._tc.get_or_add_tcPr().append(shd)
                            for para in cell.paragraphs:
                                for run in para.runs: run.font.bold = True
                    if "{남색}" in first_cell_text:
                        for cell in row_cells:
                            shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '000080'); cell._tc.get_or_add_tcPr().append(shd)
                            for para in cell.paragraphs:
                                for run in para.runs: run.font.bold = True; run.font.color.rgb = RGBColor.from_string('FFFFFF')
                        for idx in range(len(row_cells) - 1):
                            left_cell, right_cell = row_cells[idx], row_cells[idx + 1]
                            set_cell_border(left_cell, right={'val': 'single', 'sz': '4', 'color': 'FFFFFF'})
                            set_cell_border(right_cell, left={'val': 'single', 'sz': '4', 'color': 'FFFFFF'})
                if table_settings.get('borderless'):
                    border_attrs = {"val": "nil"}
                    for row in table.rows:
                        for cell in row.cells: set_cell_border(cell, top=border_attrs, bottom=border_attrs, left=border_attrs, right=border_attrs)
                merge_empty_cells_in_table(table)
                header_count = 0
                for idx in range(min(5, len(table.rows))):
                    if "{제목행}" in table.rows[idx].cells[0].text: header_count = idx + 1
                if header_count:
                    for idx in range(header_count):
                        tr = table.rows[idx]._tr
                        trPr = tr.get_or_add_trPr()
                        tblHeader = OxmlElement('w:tblHeader'); trPr.append(tblHeader)
            i += 1
            continue
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue
        p = doc.add_paragraph()
        para_format = p.paragraph_format
        if "{왼쪽}" in line: para_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif "{가운데}" in line: para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "{오른쪽}" in line: para_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif "{양쪽}" in line: para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif "{균등}" in line: para_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        indent_match = re.search(r'\{들여쓰기,1번줄:([0-9]+(?:\.[0-9]+)?),2번줄이하:([0-9]+(?:\.[0-9]+)?)\}', line)
        if indent_match:
            first_line_cm = float(indent_match.group(1))
            hanging_cm = float(indent_match.group(2))
            para_format.left_indent = Cm(hanging_cm)
            para_format.first_line_indent = Cm(first_line_cm - hanging_cm)
        line_spacing_match = re.search(r'\{(\d+(?:\.\d+)?)줄\}', line)
        if line_spacing_match:
            para_format.line_spacing = float(line_spacing_match.group(1))
        title_match = re.search(r'\{제목(\d)\.(\d)\}', line)
        image_match = re.search(r'\{그림:([^}]+)\}', line)
        if title_match:
            level, align_num = int(title_match.group(1)), int(title_match.group(2))
            font_sizes = {1: 18, 2: 16, 3: 14}; alignments = {1: WD_ALIGN_PARAGRAPH.LEFT, 2: WD_ALIGN_PARAGRAPH.CENTER, 3: WD_ALIGN_PARAGRAPH.RIGHT}
            clean_line = re.sub(r'\{[^}]+\}', '', line); run = p.add_run(clean_line); run.bold = True
            run.font.size = Pt(font_sizes.get(level, 12)); para_format.alignment = alignments.get(align_num, WD_ALIGN_PARAGRAPH.LEFT)
        elif image_match:
            image_url_or_id = image_match.group(1).strip()
            p.clear()
            insert_image_to_paragraph(p, image_url_or_id, section)
        else:
            font_size_match = re.search(r'\{([\d.]+)pt\}', line)
            tokens = re.split(r'(\{>>\}|\{<<\}|\{탭\}|\{줄바꿈\}|\{문단바꿈\})', line)
            in_bold_underline = False
            for token in tokens:
                if not token: continue
                if token == '{>>}': in_bold_underline = True
                elif token == '{<<}': in_bold_underline = False
                elif token == '{문단바꿈}': p = doc.add_paragraph()
                elif token == '{줄바꿈}': p.add_run().add_break(WD_BREAK.LINE)
                elif token == '{탭}': p.add_run('\t')
                else:
                    clean_token = re.sub(r'\{[^}]+\}', '', token)
                    run = p.add_run(clean_token)
                    if in_bold_underline: run.font.bold = True; run.font.underline = True
                    if font_size_match: run.font.size = Pt(float(font_size_match.group(1)))
        i += 1
    # 최종 후처리
    for table in doc.tables:
        for c in table._tbl.tblPr.getchildren():
            if c.tag.endswith('tblLayout'): table._tbl.tblPr.remove(c)
        tbl_layout = OxmlElement('w:tblLayout'); tbl_layout.set(qn('w:type'), 'fixed'); table._tbl.tblPr.append(tbl_layout)
        tbl_w = OxmlElement('w:tblW'); tbl_w.set(qn('w:w'), '5000'); tbl_w.set(qn('w:type'), 'pct'); table._tbl.tblPr.append(tbl_w)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.element.find(qn('w:br')) is not None and run.element.find(qn('w:br')).get(qn('w:type')) == 'page': continue
            if run.element.find(qn('w:drawing')) is not None: continue
            run.text = re.sub(r'\{[^}]+\}', '', run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.element.find(qn('w:drawing')) is not None: continue
                        run.text = re.sub(r'\{[^}]+\}', '', run.text)
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# ==============================================================================
# 3.5: Google Sheet 로깅 헬퍼 함수 (다음 단계에서 사용 예정)
# ==============================================================================
def log_to_gemini_usage_sheet(request_text, response_text, token_count):
    print(f"Logging to Gemini Usage Sheet (skipping for now): Request='{request_text}', Response='{response_text}', Tokens={token_count}")
    pass

# ==============================================================================
# 4. Flask API 엔드포인트
# ==============================================================================

@app.route("/")
def index():
    return "<h1>SaeRo AI Editor Backend is running.</h1>"


@app.route('/check-user', methods=['POST'])
def handle_check_user():
    if not SHEET_CLIENT:
        return jsonify({"error": "Google Sheets service is not available"}), 503

    try:
        if not request.is_json:
            return jsonify({"error": "Request must be JSON"}), 400

        data = request.get_json()
        name = data.get('name')
        email = data.get('email')

        if not name or not email:
            return jsonify({"error": "Name and email are required"}), 400

        # 구글 시트 열기
        spreadsheet = SHEET_CLIENT.open_by_key("10FWgDt04ox83Fc2iDM66seswL1k2W-rfhOE1GHrjZtI")
        worksheet = spreadsheet.worksheet("시트1")

        # 모든 데이터 가져오기 (헤더 포함)
        all_users = worksheet.get_all_records()

        is_authorized = False
        for user_record in all_users:
            # 시트의 헤더 이름과 정확히 일치해야 함
            if (user_record.get('사용자이름') == name and
                user_record.get('이메일') == email and
                str(user_record.get('상태')).strip() == '1'):
                is_authorized = True
                break

        return jsonify({"authorized": is_authorized})

    except gspread.exceptions.SpreadsheetNotFound:
         print("Error: Spreadsheet not found. Check the key and sharing settings.")
         return jsonify({"error": "Could not access the user list spreadsheet."}), 500
    except Exception as e:
        print(f"Error in /check-user: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500


@app.route('/create-docx', methods=['POST'])
def handle_create_docx():
    try:
        if not request.is_json: return jsonify({"error": "Missing JSON in request"}), 400
        data = request.get_json()
        if 'content' not in data or 'settings' not in data:
            return jsonify({"error": "Missing 'content' or 'settings' in request body"}), 400
        file_stream = create_word_document(data['content'], data['settings'])
        title = data.get('title', '').strip()
        filename = generate_dynamic_filename(title)
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/chat-gemini', methods=['POST'])
def handle_chat():
    if model is None:
        return jsonify({"error": "Gemini API 모델이 초기화되지 않았습니다."}), 503
    try:
        if not request.is_json:
            return jsonify({"error": "요청 형식이 올바르지 않습니다. (JSON 필요)"}), 400

        data = request.get_json()
        user_message = data.get('message')
        chat_history = data.get('history', [])

        if not user_message:
            return jsonify({"error": "'message' 필드가 요청에 포함되지 않았습니다."}), 400

    except Exception as e:
        print(f"요청 데이터 처리 중 오류 발생: {e}")
        return jsonify({"error": "요청 데이터를 파싱하는 중 오류가 발생했습니다."}), 400

    try:
        chat_session = model.start_chat(history=chat_history)
        response = chat_session.send_message(user_message)

        # 다음 단계에서 토큰 계산 및 로깅 구현 예정
        # total_tokens = response.usage_metadata.total_token_count
        # log_to_gemini_usage_sheet(user_message, response.text, total_tokens)

        return jsonify({"reply": response.text})

    except Exception as e:
        error_message = f"AI 통신 오류: {str(e)}"
        # 다음 단계에서 오류 로깅 구현 예정
        # log_to_gemini_usage_sheet(user_message, error_message, 0)

        print(f"!!! Gemini API 호출 오류: {e}")
        return jsonify({"error": error_message}), 500

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=int(os.environ.get("PORT", 8080)))
